import os
import re
from typing import Dict, Any, List

import docx
import pdfplumber

from src.ai.ai_resume_parser import (
    parse_resume_ai,
)

from src.ai.semantic_validator import (
    filter_semantic_skills,
)

HEADINGS = [
    "CAREER OBJECTIVE",
    "OBJECTIVE",
    "PROFILE",
    "PROFESSIONAL SUMMARY",
    "SUMMARY",
    "TECHNICAL SKILLS",
    "TECH SKILLS",
    "SKILLS",
    "KEY SKILLS",
    "EXPERIENCE",
    "WORK EXPERIENCE",
    "PROFESSIONAL EXPERIENCE",
    "INTERNSHIP",
    "INTERNSHIPS",
    "PROJECTS",
    "ACADEMIC PROJECTS",
    "EDUCATION",
    "ACADEMIC DETAILS",
    "ACADEMICS",
    "QUALIFICATION",
    "CERTIFICATIONS",
    "CERTIFICATES",
    "ACHIEVEMENTS",
    "AWARDS",
    "SOFT SKILLS",
    "LANGUAGES",
]


TECH_SKILLS = {
    "Python": ["python"],
    "Java": ["java"],
    "Core Java": ["core java"],
    "C": [" c ", " c,", " c\n"],
    "C++": ["c++"],
    "C#": ["c#"],
    "JavaScript": ["javascript", " js "],
    "TypeScript": ["typescript"],
    "HTML": ["html"],
    "CSS": ["css"],
    "React.js": ["react", "react.js", "reactjs"],
    "Node.js": ["node.js", "nodejs"],
    "Express.js": ["express.js", "expressjs", "express"],
    "Next.js": ["next.js", "nextjs"],
    "Tailwind CSS": ["tailwind css", "tailwind"],
    "Bootstrap": ["bootstrap"],
    "FastAPI": ["fastapi"],
    "Flask": ["flask"],
    "Django": ["django"],
    "Spring Boot": ["spring boot"],
    "REST API": ["rest api", "restful api"],
    "MongoDB": ["mongodb", "mongo db"],
    "MySQL": ["mysql"],
    "PostgreSQL": ["postgresql", "postgres"],
    "SQLite": ["sqlite"],
    "Firebase": ["firebase"],
    "Git": [" git ", "git,"],
    "GitHub": ["github"],
    "Docker": ["docker"],
    "Kubernetes": ["kubernetes", "k8s"],
    "AWS": ["aws", "amazon web services"],
    "Azure": ["azure"],
    "GCP": ["gcp", "google cloud"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    "Excel": ["excel", "ms excel"],
    "SQL": ["sql"],
    "DAX": ["dax"],
    "Data Analysis": ["data analysis", "data analyst"],
    "Data Visualization": ["data visualization", "visualization"],
    "Data Modeling": ["data modeling"],
    "Machine Learning": ["machine learning", " ml "],
    "Deep Learning": ["deep learning", " dl "],
    "Artificial Intelligence": ["artificial intelligence", " ai "],
    "NLP": ["nlp", "natural language processing"],
    "Computer Vision": ["computer vision"],
    "NumPy": ["numpy"],
    "Pandas": ["pandas"],
    "Matplotlib": ["matplotlib"],
    "Seaborn": ["seaborn"],
    "Scikit-learn": ["scikit-learn", "sklearn", "scikit learn"],
    "TensorFlow": ["tensorflow"],
    "PyTorch": ["pytorch"],
    "Keras": ["keras"],
    "Jupyter Notebook": ["jupyter notebook", "jupyter"],
    "Google Colab": ["google colab", "colab"],
    "VS Code": ["vs code", "visual studio code"],
    "Playwright": ["playwright"],
    "BeautifulSoup": ["beautifulsoup", "beautiful soup", "bs4"],
    "Selenium": ["selenium"],
    "Web Scraping": ["web scraping", "scraping"],
    "JWT": ["jwt"],
    "OAuth": ["oauth"],
    "Authentication": ["authentication"],
    "Authorization": ["authorization"],
}


SOFT_SKILLS = {
    "Communication": ["communication"],
    "Leadership": ["leadership", "team lead"],
    "Teamwork": ["teamwork", "team work"],
    "Problem Solving": ["problem solving", "problem-solving"],
    "Critical Thinking": ["critical thinking"],
    "Analytical Thinking": ["analytical thinking"],
    "Time Management": ["time management"],
    "Adaptability": ["adaptability"],
}


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def normalize_lines(text: str) -> str:
    return (text or "").replace("\r", "\n")


def extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group().strip() if match else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"(\+91[\s-]?)?[6-9]\d{9}",
        r"\+?\d{1,3}[\s-]?\d{7,14}",
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group().strip()

    return ""


def normalize_url(url: str) -> str:
    url = url.strip().rstrip(".,;|)")
    if not url:
        return ""

    if not url.startswith("http"):
        url = "https://" + url

    return url


def extract_linkedin(text: str) -> str:
    match = re.search(
        r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9\-_%./]+",
        text,
        re.I,
    )
    return normalize_url(match.group()) if match else ""


def extract_github(text: str) -> str:
    match = re.search(
        r"(https?://)?(www\.)?github\.com/[A-Za-z0-9._\-]+/?",
        text,
        re.I,
    )

    if match:
        return normalize_url(match.group())

    username_match = re.search(r"github[:\s\-]+([A-Za-z0-9._\-]+)", text, re.I)

    if username_match:
        return f"https://github.com/{username_match.group(1).strip()}"

    return ""


def extract_portfolio(text: str) -> str:
    links = re.findall(r"(https?://[^\s]+|www\.[^\s]+)", text, re.I)

    blocked = ["linkedin.com", "github.com", "mailto:", "facebook.com", "instagram.com"]

    for link in links:
        lower = link.lower()

        if not any(site in lower for site in blocked):
            return normalize_url(link)

    return ""


def extract_name(raw_text: str) -> str:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]

    bad_words = [
        "resume",
        "curriculum",
        "vitae",
        "email",
        "phone",
        "mobile",
        "linkedin",
        "github",
        "portfolio",
        "address",
        "objective",
        "summary",
    ]

    for line in lines[:10]:
        lower = line.lower()

        if any(word in lower for word in bad_words):
            continue

        if "@" in line or "http" in lower or re.search(r"\d", line):
            continue

        if len(line.split()) <= 5 and len(line) <= 60:
            return line.title()

    return ""


def get_section_lines(raw_text: str, heading: str) -> List[str]:
    text = normalize_lines(raw_text)
    lines = [line.rstrip() for line in text.splitlines()]

    start = -1

    for i, line in enumerate(lines):
        normalized = line.strip().upper().replace(":", "")
        if normalized == heading.upper():
            start = i + 1
            break

    if start == -1:
        return []

    collected = []

    for line in lines[start:]:
        clean_line = line.strip()
        normalized = clean_line.upper().replace(":", "")

        if normalized in HEADINGS:
            break

        if clean_line:
            collected.append(clean_line)

    return collected


def get_first_available_section(raw_text: str, headings: List[str]) -> List[str]:
    for heading in headings:
        lines = get_section_lines(raw_text, heading)
        if lines:
            return lines

    return []


def extract_education_details(raw_text: str) -> Dict[str, str]:
    lines = get_first_available_section(
        raw_text,
        ["EDUCATION", "ACADEMIC DETAILS", "ACADEMICS", "QUALIFICATION"],
    )

    if not lines:
        return {
            "qualification": "",
            "course": "",
            "college": "",
            "university": "",
            "education": "",
        }

    text = " ".join(lines)

    degree_patterns = [
        r"\bMCA\b",
        r"\bBCA\b",
        r"\bMBA\b",
        r"\bBBA\b",
        r"\bB\.?\s?Tech\b",
        r"\bM\.?\s?Tech\b",
        r"\bB\.?\s?Sc\b",
        r"\bM\.?\s?Sc\b",
        r"\bB\.?\s?Com\b",
        r"\bM\.?\s?Com\b",
        r"\bBachelor of [A-Za-z ]+",
        r"\bMaster of [A-Za-z ]+",
    ]

    qualification = ""

    for pattern in degree_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            qualification = match.group().strip()
            break

    course_patterns = [
        r"Computer Applications",
        r"Computer Science",
        r"Data Science",
        r"Information Technology",
        r"Artificial Intelligence",
        r"Machine Learning",
        r"Electronics",
        r"Mechanical",
        r"Civil",
        r"Commerce",
        r"Business Administration",
    ]

    course = ""

    for pattern in course_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            course = match.group().strip()
            break

    college = ""
    university = ""

    for line in lines:
        lower = line.lower()

        if any(word in lower for word in ["university", "aktu", "affiliated"]):
            university = line.strip()

        if any(
            word in lower
            for word in ["college", "institute", "school", "academy", "iimt"]
        ):
            college = line.strip()

    return {
        "qualification": qualification,
        "course": course,
        "college": college,
        "university": university,
        "education": "\n".join(lines),
    }


def extract_experience(raw_text: str) -> str:
    lines = get_first_available_section(
        raw_text,
        [
            "EXPERIENCE",
            "WORK EXPERIENCE",
            "PROFESSIONAL EXPERIENCE",
            "INTERNSHIP",
            "INTERNSHIPS",
        ],
    )
    return "\n".join(lines) if lines else ""


def extract_experience_years(text: str) -> float:
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*(?:years|year|yrs|yr)\s+(?:of\s+)?(?:experience|exp)",
        r"(?:experience|exp)\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*(?:years|year|yrs|yr)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            try:
                return float(match.group(1))
            except Exception:
                return 0

    if any(
        word in text.lower() for word in ["fresher", "student", "internship", "intern"]
    ):
        return 0

    return 0


def extract_projects(raw_text: str) -> List[str]:
    lines = get_first_available_section(raw_text, ["PROJECTS", "ACADEMIC PROJECTS"])

    if not lines:
        return []

    projects = []
    buffer = []

    for line in lines:
        if re.match(r"^[•\-*]", line) and buffer:
            projects.append(" ".join(buffer).strip())
            buffer = [line.lstrip("•-* ").strip()]
        else:
            buffer.append(line.strip())

    if buffer:
        projects.append(" ".join(buffer).strip())

    return [project for project in projects if project]


def extract_certifications(raw_text: str) -> List[str]:
    lines = get_first_available_section(raw_text, ["CERTIFICATIONS", "CERTIFICATES"])
    return lines if lines else []


def extract_achievements(raw_text: str) -> List[str]:
    lines = get_first_available_section(raw_text, ["ACHIEVEMENTS", "AWARDS"])
    return lines if lines else []


def extract_skills_from_map(text: str, skill_map: Dict[str, List[str]]) -> List[str]:
    lower = f" {text.lower()} "
    skills = []

    for skill, patterns in skill_map.items():
        if any(pattern.lower() in lower for pattern in patterns):
            skills.append(skill)

    return sorted(set(skills))


def extract_skills(text: str) -> List[str]:
    return extract_skills_from_map(text, TECH_SKILLS)


def extract_soft_skills(text: str) -> List[str]:
    return extract_skills_from_map(text, SOFT_SKILLS)


def categorize_skills(skills: List[str]) -> Dict[str, List[str]]:
    frameworks = []
    libraries = []
    databases = []
    tools = []
    domains = []

    framework_set = {
        "React.js",
        "Node.js",
        "Express.js",
        "Next.js",
        "FastAPI",
        "Flask",
        "Django",
        "Spring Boot",
        "Tailwind CSS",
        "Bootstrap",
    }

    library_set = {
        "NumPy",
        "Pandas",
        "Matplotlib",
        "Seaborn",
        "Scikit-learn",
        "TensorFlow",
        "PyTorch",
        "Keras",
    }

    database_set = {"MongoDB", "MySQL", "PostgreSQL", "SQLite", "Firebase"}

    tool_set = {
        "Git",
        "GitHub",
        "Docker",
        "Kubernetes",
        "Power BI",
        "Tableau",
        "Excel",
        "Jupyter Notebook",
        "Google Colab",
        "VS Code",
        "Playwright",
        "Selenium",
        "BeautifulSoup",
    }

    domain_set = {
        "Data Analysis",
        "Data Visualization",
        "Machine Learning",
        "Deep Learning",
        "Artificial Intelligence",
        "NLP",
        "Computer Vision",
        "Web Scraping",
        "Data Modeling",
    }

    for skill in skills:
        if skill in framework_set:
            frameworks.append(skill)
        if skill in library_set:
            libraries.append(skill)
        if skill in database_set:
            databases.append(skill)
        if skill in tool_set:
            tools.append(skill)
        if skill in domain_set:
            domains.append(skill)

    return {
        "frameworks": sorted(set(frameworks)),
        "libraries": sorted(set(libraries)),
        "databases": sorted(set(databases)),
        "tools": sorted(set(tools)),
        "domains": sorted(set(domains)),
    }


def detect_role_from_skills(skills: List[str], text: str) -> str:
    combined = (" ".join(skills) + " " + text).lower()

    if any(
        x in combined
        for x in [
            "power bi",
            "tableau",
            "excel",
            "sql",
            "data analysis",
            "pandas",
            "dashboard",
        ]
    ):
        return "Data Analyst"

    if any(
        x in combined
        for x in [
            "machine learning",
            "deep learning",
            "tensorflow",
            "pytorch",
            "scikit-learn",
            "nlp",
        ]
    ):
        return "AI/ML Engineer"

    if any(x in combined for x in ["react", "node", "mongodb", "express"]):
        return "MERN Developer"

    if any(x in combined for x in ["fastapi", "django", "flask", "python"]):
        return "Python Developer"

    if any(x in combined for x in ["html", "css", "javascript", "tailwind"]):
        return "Frontend Developer"

    return "General Candidate"


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")

        return "\n".join(text_parts)

    if ext == ".docx":
        document = docx.Document(file_path)

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        table_text = []

        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_text.append(row_text)

        return "\n".join(paragraphs + table_text)

    raise ValueError("Only PDF/DOCX allowed")


def fallback_parse(raw_text: str, clean: str) -> Dict[str, Any]:
    skills = extract_skills(clean)
    soft_skills = extract_soft_skills(clean)
    categorized = categorize_skills(skills)
    education_data = extract_education_details(raw_text)
    experience = extract_experience(raw_text)
    projects = extract_projects(raw_text)

    role = detect_role_from_skills(skills, clean)

    return {
        "name": extract_name(raw_text),
        "email": extract_email(clean),
        "phone": extract_phone(clean),
        "linkedin": extract_linkedin(raw_text),
        "github": extract_github(raw_text),
        "portfolio": extract_portfolio(raw_text),
        "location": "",
        "skills": skills,
        "technicalSkills": skills,
        "softSkills": soft_skills,
        "tools": categorized["tools"],
        "libraries": categorized["libraries"],
        "frameworks": categorized["frameworks"],
        "databases": categorized["databases"],
        "domains": categorized["domains"],
        "qualification": education_data["qualification"],
        "course": education_data["course"],
        "college": education_data["college"],
        "university": education_data["university"],
        "education": education_data["education"],
        "experience": experience,
        "experienceYears": extract_experience_years(clean),
        "projects": projects,
        "certifications": extract_certifications(raw_text),
        "achievements": extract_achievements(raw_text),
        "role": role,
        "summary": f"{role} profile extracted from resume evidence.",
        "resume_text": clean,
        "source": "fallback",
    }


def as_list(value: Any) -> List[str]:
    if value is None:
        return []

    if isinstance(value, list):
        raw = value
    elif isinstance(value, str):
        raw = [x.strip() for x in value.replace("|", ",").split(",")]
    else:
        raw = [str(value)]

    cleaned = []

    for item in raw:
        if isinstance(item, dict):
            parts = []

            for key in ["title", "name", "description", "techStack", "technologies"]:
                if item.get(key):
                    parts.append(str(item.get(key)))

            item = " - ".join(parts) if parts else str(item)

        item = str(item).strip()

        if item:
            cleaned.append(item)

    return sorted(set(cleaned))


def merge_ai_with_fallback(
    ai_data: Dict[str, Any],
    fallback_data: Dict[str, Any],
) -> Dict[str, Any]:

    final_data = fallback_data.copy()

    for key, value in ai_data.items():
        if value not in ["", None, [], {}]:
            final_data[key] = value

    protected_keys = [
        "name",
        "email",
        "phone",
        "linkedin",
        "github",
        "portfolio",
        "qualification",
        "course",
        "college",
        "university",
    ]

    for key in protected_keys:
        if fallback_data.get(key):
            final_data[key] = fallback_data[key]

    list_fields = [
        "skills",
        "technicalSkills",
        "softSkills",
        "tools",
        "libraries",
        "frameworks",
        "databases",
        "domains",
        "projects",
        "certifications",
        "achievements",
    ]

    for field in list_fields:
        final_data[field] = sorted(
            set(as_list(fallback_data.get(field)) + as_list(ai_data.get(field)))
        )

    all_skills = []

    for field in [
        "skills",
        "technicalSkills",
        "tools",
        "libraries",
        "frameworks",
        "databases",
        "domains",
    ]:
        all_skills.extend(as_list(final_data.get(field)))

    final_data["skills"] = sorted(set(all_skills))

    role = final_data.get(
        "role",
        "",
    )

    semantic_result = filter_semantic_skills(
        role=role,
        skills=final_data.get(
            "skills",
            [],
        ),
    )

    final_data["skills"] = semantic_result["valid"]

    final_data["removedSkills"] = semantic_result["removed"]

    final_data["resume_text"] = fallback_data["resume_text"]

    final_data["source"] = f"{ai_data.get('source', 'AI')}+fallback"

    return final_data


def parse_resume(file_path: str) -> Dict[str, Any]:
    raw_text = extract_text(file_path)

    if not raw_text or not raw_text.strip():
        raise ValueError(
            "Empty resume or unsupported scanned PDF. Please upload a text-based PDF/DOCX."
        )

    clean = clean_text(raw_text)

    fallback_data = fallback_parse(raw_text, clean)
    ai_data = parse_resume_ai(clean)

    if ai_data:
        return merge_ai_with_fallback(ai_data, fallback_data)

    return fallback_data
